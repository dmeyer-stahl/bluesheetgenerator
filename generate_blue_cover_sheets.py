import pdfplumber
import re
from docx import Document
from datetime import datetime
import os
import json

def load_masterformat_data(base_path):
    """Loads the MasterFormat JSON files."""
    valid_sections_path = os.path.join(base_path, "valid_sections.json")
    section_titles_path = os.path.join(base_path, "section_titles.json")

    if not os.path.exists(valid_sections_path) or not os.path.exists(section_titles_path):
        print("❌ Error: 'valid_sections.json' and 'section_titles.json' not found.")
        print(f"   Please ensure they are in the same directory as the script or PDF.")
        exit(1)

    with open(valid_sections_path, 'r') as f:
        valid_sections = set(json.load(f))
    with open(section_titles_path, 'r') as f:
        section_titles = json.load(f)
    
    return valid_sections, section_titles

def extract_sections_from_pdf(pdf_path, valid_sections, section_titles):
    """
    Extracts validated CSI MasterFormat sections from a PDF Table of Contents.
    """
    # Use a dictionary to store found sections to prevent duplicates
    extracted_sections = {}
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue

            # Stop processing if the 'DRAWINGS' section is reached
            if "DRAWINGS" in text.upper():
                break

            for line in text.splitlines():
                line = line.strip()
                # Regex to find patterns like "031000", "03 10 00", or "03-10-00" at the start of a line
                match = re.match(r'^(\d[\d\s\-]{4,8})', line)
                
                if not match:
                    continue

                # Normalize the found number (remove spaces/hyphens)
                potential_number = match.group(1)
                normalized_number = re.sub(r'[\s\-]', '', potential_number)

                # Check if the normalized number is a valid section
                if normalized_number in valid_sections:
                    # If valid, use the official title from the JSON file
                    official_title = section_titles.get(normalized_number, "Unknown Title")
                    extracted_sections[normalized_number] = official_title

    # Return the sections as a list of (number, title) tuples
    return list(extracted_sections.items())

def create_cover_sheet(doc, section_number, section_title):
    """Creates a single cover sheet page in the Word document."""
    # Format the section number with spaces for readability (e.g., 03 10 00)
    formatted_number = f"{section_number[:2]} {section_number[2:4]} {section_number[4:]}"
    
    doc.add_paragraph().add_run(f"{formatted_number} – {section_title.upper()}").bold = True
    doc.add_paragraph("") # Spacer

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Number"
    hdr_cells[1].text = "Alternates"

    # Add empty placeholder rows for notes
    doc.add_paragraph("\nNotes")
    for _ in range(8):
        doc.add_paragraph("")
    doc.add_page_break()

def create_blue_sheet_doc(sections, project_name, bid_date, output_path):
    """Creates the final Word document with all cover sheets."""
    doc = Document()
    doc.add_paragraph().add_run(project_name).bold = True
    doc.add_paragraph(f"Bid Date: {bid_date}")
    doc.add_page_break()

    # Always include a cover sheet for Final Cleaning (Division 1)
    create_cover_sheet(doc, "017400", "Final Cleaning")

    # Create a cover sheet for each extracted section
    for section_number, section_title in sorted(sections): # Sort sections numerically
        create_cover_sheet(doc, section_number, section_title)

    doc.save(output_path)
    print(f"\n✅ Blue cover sheets saved to: {output_path}")

if __name__ == "__main__":
    print("=== Blue Sheet Generator (Upgraded) ===")
    
    # --- User Input ---
    pdf_path = input("Enter path to TOC PDF file: ").strip().strip('"')
    if not os.path.exists(pdf_path):
        print(f"❌ File not found: {pdf_path}")
        exit(1)

    project_name = input("Enter project name: ").strip()
    bid_date_str = input("Enter bid date (YYYY-MM-DD): ").strip()
    try:
        # Validate date format but use the string for output
        datetime.strptime(bid_date_str, "%Y-%m-%d")
    except ValueError:
        print("❌ Invalid date format. Please use YYYY-MM-DD.")
        exit(1)
        
    # --- Processing ---
    # Assume JSON files are in the same directory as the PDF
    script_directory = os.path.dirname(pdf_path)
    valid_sections, section_titles = load_masterformat_data(script_directory)
    
    sections = extract_sections_from_pdf(pdf_path, valid_sections, section_titles)
    
    if not sections:
        print("\n⚠️ No valid sections were found in the PDF. The output will only contain the default pages.")
    else:
        print(f"\nFound {len(sections)} valid sections.")

    output_path = os.path.join(script_directory, "Generated Blue Cover Sheets.docx")
    create_blue_sheet_doc(sections, project_name, bid_date_str, output_path)