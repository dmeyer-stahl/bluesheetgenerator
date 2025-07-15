import base64
import io
import re
from datetime import datetime

from docx import Document
from flask import Flask, request, jsonify, send_file
import pdfplumber

# --- Flask App Initialization ---
app = Flask(__name__)

# --- Helper Functions (copy these from your original script) ---
# NOTE: Make sure valid_sections.json and section_titles.json are in your GitHub repo.
def load_masterformat_data(base_path):
    with open('valid_sections.json', 'r') as f:
        valid_sections = set(json.load(f))
    with open('section_titles.json', 'r') as f:
        section_titles = json.load(f)
    return valid_sections, section_titles
    
def extract_sections_from_pdf(pdf_file, valid_sections, section_titles):
    extracted_sections = {}
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            if "DRAWINGS" in text.upper():
                break
            for line in text.splitlines():
                line = line.strip()
                match = re.match(r'^(\d[\d\s\-]{4,8})', line)
                if not match:
                    continue
                potential_number = match.group(1)
                normalized_number = re.sub(r'[\s\-]', '', potential_number)
                if normalized_number in valid_sections:
                    official_title = section_titles.get(normalized_number, "Unknown Title")
                    extracted_sections[normalized_number] = official_title
    return list(extracted_sections.items())

def create_cover_sheet(doc, section_number, section_title):
    formatted_number = f"{section_number[:2]} {section_number[2:4]} {section_number[4:]}"
    doc.add_paragraph().add_run(f"{formatted_number} – {section_title.upper()}").bold = True
    doc.add_paragraph("")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Number"
    hdr_cells[1].text = "Alternates"
    doc.add_paragraph("\nNotes")
    for _ in range(8):
        doc.add_paragraph("")
    doc.add_page_break()

# --- API Endpoint Definition ---
@app.route('/generate', methods=['POST'])
def generate_blue_sheets():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON"}), 400

        project_name = data['projectName']
        bid_date_str = data['bidDate']
        pdf_b64 = data['pdfBase64']
        
        pdf_bytes = base64.b64decode(pdf_b64)
        pdf_file = io.BytesIO(pdf_bytes)

    except (KeyError, TypeError) as e:
        return jsonify({"error": f"Missing required field: {e}"}), 400

    # Load masterformat data
    valid_sections, section_titles = load_masterformat_data('.')

    # Run the core logic
    sections = extract_sections_from_pdf(pdf_file, valid_sections, section_titles)
    
    # --- Create the DOCX in memory ---
    doc = Document()
    doc_stream = io.BytesIO()

    doc.add_paragraph().add_run(project_name).bold = True
    doc.add_paragraph(f"Bid Date: {bid_date_str}")
    doc.add_page_break()
    create_cover_sheet(doc, "017400", "Final Cleaning")
    for section_number, section_title in sorted(sections):
        create_cover_sheet(doc, section_number, section_title)

    doc.save(doc_stream)
    doc_stream.seek(0)

    return send_file(
        doc_stream,
        as_attachment=True,
        download_name=f'{project_name}_Blue_Sheets.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    # This part is for local testing only
    app.run(debug=True)
